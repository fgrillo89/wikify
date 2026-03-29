"""Export generated markdown to DOCX with journal-profile styling."""

from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from rich.console import Console

from scholarforge.export.chemistry import split_formula_runs
from scholarforge.export.journal_profile import JournalProfile
from scholarforge.store.models import Paper

console = Console()

_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[\d+\])")

# Detects:  ![Figure N: short caption](figure_N_placeholder.png)
_FIGURE_IMG_RE = re.compile(r"!\[Figure\s+(\d+):[^\]]*\]", re.IGNORECASE)
# Detects:  **Figure N.** Detailed caption text ...
_FIGURE_CAPTION_RE = re.compile(r"\*\*Figure\s+\d+\.\*\*", re.IGNORECASE)


def _parse_inline(
    text: str, *, superscript_citations: bool = True
) -> list[tuple[str, bool, bool, bool]]:
    """Return (text, bold, italic, superscript) tuples for a line."""
    parts: list[tuple[str, bool, bool, bool]] = []
    for segment in _INLINE_RE.split(text):
        if not segment:
            continue
        if segment.startswith("**") and segment.endswith("**"):
            parts.append((segment[2:-2], True, False, False))
        elif segment.startswith("*") and segment.endswith("*"):
            parts.append((segment[1:-1], False, True, False))
        elif re.fullmatch(r"\[\d+\]", segment):
            if superscript_citations:
                parts.append((segment[1:-1], False, False, True))
            else:
                parts.append((segment, False, False, False))
        else:
            parts.append((segment, False, False, False))
    return parts


def _set_run_font(run, font_family: str, font_size_pt: int | None = None) -> None:
    """Apply font family (and optional size) to a run."""
    run.font.name = font_family
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_family)
    r_fonts.set(qn("w:hAnsi"), font_family)
    r_fonts.set(qn("w:cs"), font_family)
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)


def _set_superscript(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    vert_align = OxmlElement("w:vertAlign")
    vert_align.set(qn("w:val"), "superscript")
    r_pr.append(vert_align)


def _set_subscript(run) -> None:
    r_pr = run._r.get_or_add_rPr()
    vert_align = OxmlElement("w:vertAlign")
    vert_align.set(qn("w:val"), "subscript")
    r_pr.append(vert_align)


class DocxExporter:
    """Export a numbered-markdown string to a styled DOCX file."""

    def __init__(self, journal_profile: JournalProfile) -> None:
        self.profile = journal_profile
        self._using_template = False
        self._exemplars: dict[str, OxmlElement] = {}

    def export(
        self,
        numbered_markdown: str,
        ordered_papers: list[Paper],  # noqa: ARG002
        output_path: Path,
    ) -> Path:
        from scholarforge.export.templates.registry import get_template_path

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        template = get_template_path(self.profile.template_docx)
        if template:
            doc = Document(str(template))
            self._using_template = True
            self._collect_exemplars(doc)
            self._clear_body(doc)
            console.print(f"[dim]Using template: {template.name}[/dim]")
        else:
            doc = Document()
            self._using_template = False
            self._configure_document(doc)

        self._parse_markdown(doc, numbered_markdown)

        doc.save(str(output_path))
        console.print(f"[green]DOCX exported:[/green] {output_path}")
        return output_path

    # ------------------------------------------------------------------
    # Template handling
    # ------------------------------------------------------------------

    def _collect_exemplars(self, doc: Document) -> None:
        """Save a deep copy of one paragraph per style from the template.

        These exemplars preserve the exact XML formatting (spacing, font,
        indentation) that the template defines. When we need a new paragraph
        of a given style, we clone from the exemplar instead of using
        doc.add_paragraph() which may not inherit all XML properties.
        """
        self._exemplars = {}
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            if style_name not in self._exemplars:
                # Deep-copy the paragraph XML, strip all runs (text content)
                p_copy = copy.deepcopy(para._element)
                for r in p_copy.findall(qn("w:r")):
                    p_copy.remove(r)
                self._exemplars[style_name] = p_copy

    def _add_para_from_exemplar(self, doc: Document, style_name: str):
        """Clone a paragraph from the template exemplar, preserving formatting."""
        from docx.text.paragraph import Paragraph

        if style_name in self._exemplars:
            new_p = copy.deepcopy(self._exemplars[style_name])
            # Insert before the last body-level sectPr
            body = doc.element.body
            sect_pr = body.find(qn("w:sectPr"))
            if sect_pr is not None:
                sect_pr.addprevious(new_p)
            else:
                body.append(new_p)
            return Paragraph(new_p, doc.element.body)

        # Fallback: use add_paragraph
        return doc.add_paragraph(style=style_name)

    @staticmethod
    def _clear_body(doc: Document) -> None:
        """Remove placeholder content, preserving section breaks."""
        body = doc.element.body
        sect_pr_tag = qn("w:sectPr")
        for child in list(body):
            if child.tag.endswith("}tbl"):
                body.remove(child)
            elif child.tag.endswith("}p"):
                if child.find(f".//{sect_pr_tag}") is not None:
                    for r in child.findall(qn("w:r")):
                        child.remove(r)
                else:
                    body.remove(child)

    def _configure_document(self, doc: Document) -> None:
        """Set page margins and default styles (blank document only)."""
        for section in doc.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        normal = doc.styles["Normal"]
        normal.font.name = self.profile.font_family
        normal.font.size = Pt(self.profile.font_size_pt)
        normal.paragraph_format.line_spacing = self.profile.line_spacing

        for level in range(1, 4):
            try:
                doc.styles[f"Heading {level}"].font.name = self.profile.font_family
            except KeyError:
                pass

    def _style_name(self, role: str) -> str:
        return self.profile.style_map.get(role, role)

    # ------------------------------------------------------------------
    # Markdown parser
    # ------------------------------------------------------------------

    def _parse_markdown(self, doc: Document, markdown: str) -> None:
        lines = markdown.splitlines()
        in_references = False
        in_abstract = False
        title_written = False

        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            # --- Figure placeholder box (![Figure N: ...] line) --------------
            if _FIGURE_IMG_RE.match(stripped):
                # Extract the short caption from the alt text for the box label
                m = re.match(r"!\[([^\]]*)\]", stripped)
                alt_text = m.group(1) if m else "Figure"
                self._add_figure_placeholder(doc, alt_text)
                continue

            # --- Figure caption line (**Figure N.** ...) ----------------------
            if _FIGURE_CAPTION_RE.match(stripped):
                # Render as an italic caption paragraph
                caption_text = stripped
                self._add_figure_caption(doc, caption_text)
                continue

            if stripped.startswith("### "):
                self._add_heading(doc, stripped[4:], level=2)
                in_abstract = False
            elif stripped.startswith("## "):
                heading_text = stripped[3:]
                self._add_heading(doc, heading_text, level=1)
                in_references = heading_text.strip().lower() == "references"
                in_abstract = False
            elif stripped.startswith("# "):
                heading_text = stripped[2:]
                in_abstract = False
                in_references = False
                if heading_text.strip().lower() == "abstract":
                    in_abstract = True
                elif not title_written:
                    self._add_title(doc, heading_text)
                    title_written = True
                else:
                    self._add_heading(doc, heading_text, level=1)
            else:
                if in_references:
                    role = "references"
                elif in_abstract:
                    role = "abstract"
                else:
                    role = "body"
                self._add_body_paragraph(
                    doc,
                    stripped,
                    superscript_citations=not in_references,
                    style_role=role,
                )

    # ------------------------------------------------------------------
    # Element builders
    # ------------------------------------------------------------------

    def _add_title(self, doc: Document, text: str) -> None:
        style = self._style_name("title")
        if self._using_template:
            para = self._add_para_from_exemplar(doc, style)
        else:
            try:
                para = doc.add_paragraph(style=style)
            except KeyError:
                para = doc.add_heading(level=0)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._fill_runs(para, text)

    def _add_heading(self, doc: Document, text: str, level: int) -> None:
        style = self._style_name(f"heading{level}")
        if self._using_template:
            para = self._add_para_from_exemplar(doc, style)
        else:
            try:
                para = doc.add_paragraph(style=style)
            except KeyError:
                para = doc.add_heading(level=level)
        self._fill_runs(para, text)

    def _add_body_paragraph(
        self,
        doc: Document,
        text: str,
        *,
        superscript_citations: bool = True,
        style_role: str = "body",
    ) -> None:
        style = self._style_name(style_role)
        if self._using_template:
            para = self._add_para_from_exemplar(doc, style)
        else:
            try:
                para = doc.add_paragraph(style=style)
            except KeyError:
                para = doc.add_paragraph(style="Normal")
            para.paragraph_format.line_spacing = self.profile.line_spacing
        self._fill_runs(para, text, superscript_citations=superscript_citations)

    def _fill_runs(self, para, text: str, *, superscript_citations: bool = True) -> None:
        tokens = _parse_inline(text, superscript_citations=superscript_citations)
        font = self.profile.font_family
        size = self.profile.font_size_pt if not self._using_template else None
        for content, bold, italic, superscript in tokens:
            if superscript or bold or italic:
                run = para.add_run(content)
                run.bold = bold
                run.italic = italic
                if not self._using_template:
                    _set_run_font(run, font, size)
                if superscript:
                    _set_superscript(run)
            else:
                self._fill_with_chemistry(para, content, font, size)

    def _fill_with_chemistry(
        self,
        para,
        text: str,
        font: str,
        size: int | None,
    ) -> None:
        words = re.split(r"(\b\w+\b)", text)
        for word in words:
            formula_runs = split_formula_runs(word)
            if len(formula_runs) == 1 and not formula_runs[0][1]:
                run = para.add_run(word)
                if not self._using_template:
                    _set_run_font(run, font, size)
            else:
                for part, is_subscript in formula_runs:
                    run = para.add_run(part)
                    if not self._using_template:
                        _set_run_font(run, font, size)
                    if is_subscript:
                        _set_subscript(run)

    # ------------------------------------------------------------------
    # Figure placeholder rendering
    # ------------------------------------------------------------------

    def _add_figure_placeholder(self, doc: Document, alt_text: str) -> None:
        """Insert a gray bordered box as a figure placeholder in the DOCX."""
        # A 1-cell table gives us reliable borders and shading
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"

        cell = table.cell(0, 0)

        # Set a fixed height via cell properties
        tc = cell._tc
        tcp = tc.get_or_add_tcPr()

        # Gray shading
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")  # light gray
        tcp.append(shd)

        # Set cell height (approx 2 inches = 2880 twips)
        tr = table.rows[0]._tr
        tr_pr = tr.get_or_add_trPr()
        tr_height = OxmlElement("w:trHeight")
        tr_height.set(qn("w:val"), "2880")
        tr_height.set(qn("w:hRule"), "exact")
        tr_pr.append(tr_height)

        # Center the placeholder text inside the cell
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(24)

        run = para.add_run("INSERT FIGURE HERE")
        run.bold = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        if not self._using_template:
            _set_run_font(run, self.profile.font_family, self.profile.font_size_pt)

        # Add a small note with the alt text below the main label
        if alt_text:
            note_para = cell.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = note_para.add_run(alt_text)
            note_run.italic = True
            note_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            if not self._using_template:
                _set_run_font(note_run, self.profile.font_family, self.profile.font_size_pt - 1)

        # Add spacing after the table
        doc.add_paragraph()

    def _add_figure_caption(self, doc: Document, text: str) -> None:
        """Add a figure caption paragraph (italic, small, centered)."""
        style = self._style_name("body")
        if self._using_template:
            para = self._add_para_from_exemplar(doc, style)
        else:
            try:
                para = doc.add_paragraph(style=style)
            except KeyError:
                para = doc.add_paragraph(style="Normal")
            para.paragraph_format.line_spacing = self.profile.line_spacing

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Strip outer ** markers and parse remainder with inline parser
        # but render all text as italic
        font = self.profile.font_family
        size = (self.profile.font_size_pt - 1) if not self._using_template else None
        tokens = _parse_inline(text, superscript_citations=False)
        for content, bold, _italic, superscript in tokens:
            run = para.add_run(content)
            run.bold = bold
            run.italic = True  # all caption text is italic
            if not self._using_template:
                _set_run_font(run, font, size)
            if superscript:
                _set_superscript(run)

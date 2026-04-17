"""DOCX parser.

Returns a ``ParseResult`` with typed ``RawImage`` records.
"""

import re
from pathlib import Path

from ..metadata import (
    choose_document_title,
    extract_document_doi,
    extract_publication_fields,
    extract_summary,
    parse_authors,
    parse_filename,
    validate_authors_against_filename,
)
from ._sections import section_spans
from .registry import ParseResult, RawImage


def parse(path: Path) -> ParseResult:
    from docx import Document

    doc = Document(str(path))
    md_text = _docx_to_markdown(doc)
    metadata = _extract_docx_metadata(doc, md_text, path.name)
    images_raw = _extract_images(doc)

    title = metadata.get("title") or path.stem
    return ParseResult(
        markdown=md_text,
        sections=section_spans(md_text),
        raw_images=images_raw,
        metadata=metadata,
        title=title,
    )


def _paragraph_to_md(para) -> str:
    style_name = para.style.name or ""
    m = re.match(r"Heading\s+(\d+)", style_name)
    if m:
        level = int(m.group(1))
        return f"{'#' * level} {para.text.strip()}"

    parts: list[str] = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            t = f"***{t}***"
        elif run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        parts.append(t)
    line = "".join(parts).strip()
    if not line:
        return ""
    if style_name.startswith("List Bullet"):
        return f"- {line}"
    if style_name.startswith("List Number"):
        return f"1. {line}"
    return line


def _table_to_md(table) -> str:
    rows = table.rows
    if not rows:
        return ""
    md_rows: list[str] = []
    for i, row in enumerate(rows):
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            md_rows.append("| " + " | ".join("---" for _ in cells) + " |")
    return "\n".join(md_rows)


def _docx_to_markdown(doc) -> str:
    from docx.oxml.ns import qn

    body = doc.element.body
    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    blocks: list[str] = []
    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            para = para_map.get(id(child))
            if para is None:
                continue
            md = _paragraph_to_md(para)
            if md:
                blocks.append(md)
        elif tag == qn("w:tbl"):
            table = table_map.get(id(child))
            if table is None:
                continue
            md = _table_to_md(table)
            if md:
                blocks.append(md)

    out: list[str] = []
    for i, b in enumerate(blocks):
        if b.startswith("#") and i > 0:
            out.append("")
        out.append(b)
    return "\n\n".join(out)


def _extract_docx_metadata(doc, md_text: str, filename: str) -> dict:
    props = doc.core_properties
    fn_year, fn_author, fn_title = parse_filename(filename)

    # Filename-first title priority: `[YYYY Author] Real Title.docx` is
    # user-curated and authoritative. core_properties.title on Word-saved
    # documents is frequently the literal "Word Document" placeholder.
    # choose_document_title walks filename > first_heading > stem and rejects
    # junk at each step, so we don't special-case "Word Document" here.
    title = choose_document_title(md_text, Path(filename))

    cp_author = (props.author or "").strip()
    if cp_author:
        authors = parse_authors(cp_author)
    elif fn_author:
        authors = [fn_author]
    else:
        authors = []
    authors = validate_authors_against_filename(authors, fn_author)
    if not authors and fn_author:
        authors = [fn_author]

    year: int | None = None
    if props.created is not None:
        try:
            year = props.created.year
        except Exception:
            year = None
    if not year and fn_year:
        year = fn_year

    metadata = {
        "title": title,
        "authors": authors,
        "summary": extract_summary(md_text),
        "year": year,
        "doi": extract_document_doi(md_text),
    }
    metadata.update(extract_publication_fields(md_text))
    return metadata


def _extract_images(doc) -> list[RawImage]:
    raw: list[RawImage] = []
    try:
        part = doc.part
        for rel in part.related_parts.values():
            content_type = getattr(rel, "content_type", "") or ""
            if not content_type.startswith("image/"):
                continue
            ext = content_type.split("/", 1)[1].split("+", 1)[0] or "png"
            blob = getattr(rel, "blob", None) or getattr(rel, "_blob", None)
            if not blob:
                continue
            raw.append(RawImage(data=blob, ext=ext))
    except Exception:
        return raw
    return raw

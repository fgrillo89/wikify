"""DOCX ingestion — converts to markdown and reuses the PDF pipeline."""

from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from rich.console import Console

from wikify.extract.chunker import chunk_sections
from wikify.extract.citations import extract_citations
from wikify.extract.figure_refs import extract_figure_refs
from wikify.extract.metadata import (
    _extract_doi,
    _extract_summary,
    _first_heading,
    _parse_authors,
    _parse_filename,
)
from wikify.ingest.pdf import ParsedPaper, _parse_section_tree, persist_parsed
from wikify.store.models import Paper

console = Console()


# ── Markdown conversion ───────────────────────────────────────────────────────


def _paragraph_to_md(para: Paragraph) -> str:
    """Convert a single docx paragraph to a markdown string."""
    style_name: str = para.style.name or ""

    # Heading → # / ## / ...
    heading_match = re.match(r"Heading\s+(\d+)", style_name)
    if heading_match:
        level = int(heading_match.group(1))
        prefix = "#" * level
        return f"{prefix} {para.text.strip()}"

    # Build text run-by-run to preserve bold/italic
    text_parts: list[str] = []
    for run in para.runs:
        run_text = run.text
        if not run_text:
            continue
        is_bold = bool(run.bold)
        is_italic = bool(run.italic)
        if is_bold and is_italic:
            run_text = f"***{run_text}***"
        elif is_bold:
            run_text = f"**{run_text}**"
        elif is_italic:
            run_text = f"*{run_text}*"
        text_parts.append(run_text)

    line = "".join(text_parts).strip()
    if not line:
        return ""

    # Bullet list
    if style_name.startswith("List Bullet"):
        return f"- {line}"

    # Numbered list (tracked externally if needed; emit as bullet for simplicity)
    if style_name.startswith("List Number"):
        return f"1. {line}"

    return line


def _table_to_md(table: Table) -> str:
    """Convert a docx Table to a simple markdown table."""
    rows = table.rows
    if not rows:
        return ""

    md_rows: list[str] = []
    for i, row in enumerate(rows):
        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
        md_rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Header separator
            md_rows.append("| " + " | ".join("---" for _ in cells) + " |")

    return "\n".join(md_rows)


def _docx_to_markdown(path: Path) -> str:
    """Convert a DOCX file to a markdown string.

    Handles headings, bold/italic runs, bullet/numbered lists, and tables.
    Paragraphs are separated by blank lines.
    """
    doc = Document(str(path))

    # Build an ordered list of block elements. python-docx exposes paragraphs
    # and tables separately; we need to reconstruct document order via the
    # raw XML body children.
    from docx.oxml.ns import qn

    body = doc.element.body

    # Map element id → object for quick lookup
    para_map = {id(p._element): p for p in doc.paragraphs}
    table_map = {id(t._element): t for t in doc.tables}

    blocks: list[str] = []

    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            para = para_map.get(id(child))
            if para is None:
                continue
            md = _paragraph_to_md(para)
            if md:
                blocks.append(md)
        elif tag == qn("w:tbl"):
            table = table_map.get(id(child))
            if table is None:
                continue
            md = _table_to_md(table)
            if md:
                blocks.append(md)

    # Join blocks: headings get a blank line before them for clean markdown
    output_parts: list[str] = []
    for i, block in enumerate(blocks):
        if block.startswith("#") and i > 0:
            output_parts.append("")  # blank line before heading
        output_parts.append(block)

    return "\n\n".join(output_parts)


# ── Metadata extraction ───────────────────────────────────────────────────────


def _extract_docx_metadata(doc: Document, md_text: str, filename: str) -> dict:
    """Build a metadata dict from DOCX core properties + markdown fallbacks."""
    props = doc.core_properties

    fn_year, fn_author, fn_title = _parse_filename(filename)

    # Title: core_properties → first heading → filename
    cp_title = (props.title or "").strip()
    heading_title = _first_heading(md_text)
    if cp_title:
        title = cp_title
    elif heading_title:
        title = heading_title
    elif fn_title:
        title = fn_title
    else:
        title = Path(filename).stem

    # Authors: core_properties.author (comma/semicolon delimited) → filename
    cp_author = (props.author or "").strip()
    if cp_author:
        authors = _parse_authors(cp_author)
    elif fn_author:
        authors = [fn_author]
    else:
        authors = []

    # Summary: scan markdown text for abstract/summary section
    summary = _extract_summary(md_text)

    # Year: core_properties.created → filename
    year: int | None = None
    if props.created is not None:
        year = props.created.year
    if not year and fn_year:
        year = fn_year

    # DOI: search in first 3000 characters of markdown
    doi = _extract_doi(md_text[:3000])

    return {
        "title": title,
        "authors": authors,
        "summary": summary,
        "year": year,
        "doi": doi,
    }


# ── Parse / ingest ────────────────────────────────────────────────────────────


def parse_docx(path: Path) -> ParsedPaper:
    """Parse a DOCX file into structured data. Does NOT touch the database or vault."""
    file_bytes = path.read_bytes()
    file_hash = hashlib.sha256(file_bytes).hexdigest()

    doc = Document(str(path))

    # Convert to markdown
    md_text = _docx_to_markdown(path)

    # Extract metadata from core properties + markdown
    metadata = _extract_docx_metadata(doc, md_text, path.name)

    # Build section tree
    section_tree = _parse_section_tree(md_text)

    # Create paper record
    paper = Paper(
        id=file_hash,
        title=metadata.get("title", path.stem),
        authors=json.dumps(metadata.get("authors", [])),
        summary=metadata.get("summary"),
        year=metadata.get("year"),
        doi=metadata.get("doi"),
        doc_type="paper",
        source_path=str(path),
        file_hash=file_hash,
        section_tree=json.dumps(section_tree),
    )

    # Chunk
    chunks = chunk_sections(md_text, section_tree, paper.id)

    # Citations from bibliography section
    citations = extract_citations(md_text, paper.id)

    # Caption-first figure references
    fig_refs = extract_figure_refs(md_text, paper.id)

    return ParsedPaper(
        paper=paper,
        chunks=chunks,
        figures=[],  # No binary figure extraction for DOCX in MVP
        citations=citations,
        figure_refs=fig_refs,
        md_text=md_text,
    )


def ingest_docx(path: Path, return_id: bool = False) -> int | str | None:
    """Ingest a single DOCX file into the knowledge base.

    Returns 1/0 or the paper ID string if return_id=True (None on skip).
    """
    file_bytes = path.read_bytes()
    file_hash = hashlib.sha256(file_bytes).hexdigest()

    from wikify.store.db import get_session

    with get_session() as session:
        existing = session.get(Paper, file_hash)
        if existing:
            console.print(f"[dim]Skipping (unchanged):[/dim] {path.name}")
            return None if return_id else 0

    parsed = parse_docx(path)
    persist_parsed(parsed)

    console.print(
        f"[green]Ingested:[/green] {path.name} "
        f"({len(parsed.chunks)} chunks, {len(parsed.figure_refs)} figure refs)"
    )
    return parsed.paper.id if return_id else 1

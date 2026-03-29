"""Create DOCX templates matching publisher formatting guidelines.

Since publisher templates are often behind Cloudflare or require manual
download, this script creates properly-formatted .docx template files
based on publisher style specifications. These serve as reference documents
with the correct styles, margins, fonts, and page setup.

Users can also download official templates manually and import them via:
    scholarforge templates import <path.docx> --name <journal_name>
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

TEMPLATES_DIR = (
    Path(__file__).parent.parent / "src" / "scholarforge" / "export" / "templates" / "docx"
)


def _set_style_font(doc, style_name: str, font_name: str, size_pt: int, **kwargs):
    """Configure a style's font properties."""
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    if "bold" in kwargs:
        style.font.bold = kwargs["bold"]
    if "color" in kwargs:
        style.font.color.rgb = kwargs["color"]


def create_wiley_afm():
    """Create a Wiley Advanced Functional Materials template.

    Based on Wiley VCH author guidelines:
    - Times New Roman, 12pt, double-spaced
    - 1-inch margins
    - Single column
    - Sections: Title, Authors, Abstract, Introduction, Results and Discussion,
      Conclusion, Experimental Section, Acknowledgements, References
    """
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Normal style (body text)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    # Title style
    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(16)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    # Heading styles
    for level, size in [(1, 14), (2, 13), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    # Save — body is empty (template only defines styles)
    path = TEMPLATES_DIR / "wiley_adv_funct_mater.docx"
    doc.save(str(path))
    print(f"Created: {path}")


def create_nature():
    """Create a Nature/Springer Nature template.

    - Times New Roman, 12pt, double-spaced
    - 5000 word limit, 150-word abstract
    - Sections: Abstract, Introduction, Results, Discussion, Methods
    """
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(18)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for level, size in [(1, 14), (2, 12), (3, 11)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    path = TEMPLATES_DIR / "nature_manuscript.docx"
    doc.save(str(path))
    print(f"Created: {path}")


def create_acs():
    """Create an ACS (American Chemical Society) template.

    - Times New Roman, 12pt, double-spaced
    - Double-column in published form, single for submission
    """
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(14)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for level, size in [(1, 12), (2, 12), (3, 11)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    path = TEMPLATES_DIR / "acs_manuscript.docx"
    doc.save(str(path))
    print(f"Created: {path}")


def create_arxiv():
    """Create an arXiv preprint template.

    - Computer Modern (fallback to Times New Roman in DOCX)
    - 11pt, ~1.15 line spacing
    - Single column, generous margins
    """
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"  # CM not available in Word
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(17)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for level, size in [(1, 14), (2, 12), (3, 11)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    path = TEMPLATES_DIR / "arxiv_preprint.docx"
    doc.save(str(path))
    print(f"Created: {path}")


def create_ieee():
    """Create an IEEE template.

    - Times New Roman, 10pt, single-spaced
    - Double column in published form
    """
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.625)
        section.right_margin = Inches(0.625)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.0

    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(24)
    title.font.bold = False  # IEEE titles are not bold
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for level, size, is_upper in [(1, 10, True), (2, 10, False), (3, 10, False)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    path = TEMPLATES_DIR / "ieee_manuscript.docx"
    doc.save(str(path))
    print(f"Created: {path}")


if __name__ == "__main__":
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    create_wiley_afm()
    create_nature()
    create_acs()
    create_arxiv()
    create_ieee()
    print(f"\nAll templates created in {TEMPLATES_DIR}")
